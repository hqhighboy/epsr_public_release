#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""从实验输出自动生成论文工件到 needs0301。"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import os
import subprocess
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List

import numpy as np
from docx import Document

from scripts.stat_tests import run_stat_tests


def _ensure(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def _read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def _read_jsonl(path: Path) -> List[Dict[str, Any]]:
    rows = []
    with path.open("r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line:
                rows.append(json.loads(line))
    return rows


def _sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        while True:
            b = f.read(8192)
            if not b:
                break
            h.update(b)
    return h.hexdigest()


def _save_plot(path_png: Path, path_pdf: Path, dpi: int = 600) -> None:
    import matplotlib.pyplot as plt

    plt.tight_layout()
    plt.savefig(path_png, dpi=dpi, bbox_inches="tight")
    plt.savefig(path_pdf, bbox_inches="tight")
    plt.close()


def _copy_file(src: Path, dst: Path) -> None:
    dst.write_bytes(src.read_bytes())


def _update_mpl_style() -> None:
    import matplotlib.pyplot as plt

    plt.rcParams.update(
        {
            "font.size": 10,
            "axes.labelsize": 11,
            "axes.titlesize": 12,
            "legend.fontsize": 9,
            "xtick.labelsize": 9,
            "ytick.labelsize": 9,
            "axes.linewidth": 1.0,
            "lines.linewidth": 1.8,
            "savefig.dpi": 600,
        }
    )


def _get_git_commit() -> str | None:
    try:
        out = subprocess.check_output(["git", "rev-parse", "HEAD"], stderr=subprocess.STDOUT)
        return out.decode("utf-8", errors="ignore").strip() or None
    except Exception:
        return None


def _filter(rows: List[Dict[str, Any]], system: str = "", algorithm: str = "") -> List[Dict[str, Any]]:
    out = rows
    if system:
        out = [r for r in out if r["system"] == system]
    if algorithm:
        out = [r for r in out if r["algorithm"] == algorithm]
    return out


def _group_systems(rows: List[Dict[str, Any]]) -> List[str]:
    return sorted({r["system"] for r in rows})


def _group_algorithms(rows: List[Dict[str, Any]]) -> List[str]:
    return sorted({r["algorithm"] for r in rows})


def create_boxplot_data(rows: List[Dict[str, Any]], out_csv: Path) -> None:
    keys = ["system", "algorithm", "run_id", "fitness", "loss_kW", "voltage_dev_pu", "cost_abs_usd_per_year"]
    with out_csv.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=keys)
        writer.writeheader()
        for r in rows:
            writer.writerow({k: r[k] for k in keys})


def create_convergence_json(rows: List[Dict[str, Any]], out_json: Path) -> None:
    payload = []
    for r in rows:
        payload.append(
            {
                "system": r["system"],
                "algorithm": r["algorithm"],
                "run_id": r["run_id"],
                "convergence_curve": r.get("convergence_curve", []),
            }
        )
    out_json.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")


def create_power_loss_fixed(rows: List[Dict[str, Any]], meta: Dict[str, Any], out_csv: Path) -> None:
    systems = _group_systems(rows)
    algorithms = _group_algorithms(rows)

    with out_csv.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["system", "item", "power_loss_kW", "unit"])
        for s in systems:
            base = meta["systems_meta"][s]["loss_base_kW"]
            writer.writerow([s, "baseline", f"{base:.6f}", "kW"])
            for a in algorithms:
                vals = [float(r["loss_kW"]) for r in rows if r["system"] == s and r["algorithm"] == a]
                if vals:
                    writer.writerow([s, a, f"{np.mean(vals):.6f}", "kW"])


def create_voltage_profile_fixed(rows: List[Dict[str, Any]], meta: Dict[str, Any], out_csv: Path) -> None:
    systems = _group_systems(rows)

    with out_csv.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["system", "node", "series", "voltage_pu", "unit"])
        for s in systems:
            base = meta["systems_meta"][s]["baseline_voltage_profile_pu"]
            for i, v in enumerate(base, start=1):
                writer.writerow([s, i, "baseline", f"{float(v):.8f}", "p.u."])

            # 每系统取 GA-PSO 最优 run 的电压曲线
            cand = [r for r in rows if r["system"] == s and r["algorithm"] == "GA-PSO"]
            if not cand:
                cand = [r for r in rows if r["system"] == s]
            if not cand:
                continue
            best = min(cand, key=lambda x: float(x["fitness"]))
            vp = best.get("voltage_profile_pu", [])
            for i, v in enumerate(vp, start=1):
                writer.writerow([s, i, f"best_{best['algorithm']}", f"{float(v):.8f}", "p.u."])


def fig_convergence(rows: List[Dict[str, Any]], system: str, out_dir: Path) -> None:
    import matplotlib.pyplot as plt

    _update_mpl_style()

    sys_rows = [r for r in rows if r["system"] == system]
    if not sys_rows:
        return
    algos = sorted({r["algorithm"] for r in sys_rows})

    plt.figure(figsize=(8, 5))
    for a in algos:
        curves = [r.get("convergence_curve", []) for r in sys_rows if r["algorithm"] == a]
        if not curves:
            continue
        max_len = max(len(c) for c in curves)
        arr = np.full((len(curves), max_len), np.nan)
        for i, c in enumerate(curves):
            arr[i, : len(c)] = np.array(c, dtype=float)
            if len(c) < max_len:
                arr[i, len(c) :] = float(c[-1])
        mean_curve = np.nanmean(arr, axis=0)
        plt.plot(np.arange(1, len(mean_curve) + 1), mean_curve, label=a, linewidth=1.8)
    plt.xlabel("Iteration")
    plt.ylabel("Fitness")
    plt.title(f"Convergence - {system}")
    plt.legend()
    plt.grid(alpha=0.3)

    if system == "IEEE33":
        png = out_dir / "Fig_Convergence_IEEE33.png"
        pdf = out_dir / "Fig_Convergence_IEEE33.pdf"
    elif system == "IEEE69":
        png = out_dir / "Fig_Convergence_Extended_IEEE69.png"
        pdf = out_dir / "Fig_Convergence_Extended_IEEE69.pdf"
    else:
        png = out_dir / "Fig_Convergence_BaituF8.png"
        pdf = out_dir / "Fig_Convergence_BaituF8.pdf"
    _save_plot(png, pdf)


def fig_boxplot_metric(rows: List[Dict[str, Any]], system: str, metric: str, title: str, filename_base: str, out_dir: Path) -> None:
    import matplotlib.pyplot as plt

    _update_mpl_style()

    sys_rows = [r for r in rows if r["system"] == system]
    if not sys_rows:
        return
    algos = sorted({r["algorithm"] for r in sys_rows})
    data = [[float(r[metric]) for r in sys_rows if r["algorithm"] == a] for a in algos]

    plt.figure(figsize=(8, 5))
    plt.boxplot(data, tick_labels=algos, showfliers=True)
    plt.title(title)
    plt.grid(axis="y", alpha=0.3)
    plt.xticks(rotation=15)

    _save_plot(out_dir / f"{filename_base}.png", out_dir / f"{filename_base}.pdf")


def fig_power_loss_compare(rows: List[Dict[str, Any]], meta: Dict[str, Any], system: str, out_dir: Path) -> None:
    import matplotlib.pyplot as plt

    _update_mpl_style()

    sys_rows = [r for r in rows if r["system"] == system]
    if not sys_rows:
        return
    algos = sorted({r["algorithm"] for r in sys_rows})
    base = float(meta["systems_meta"][system]["loss_base_kW"])
    labels = ["baseline"] + algos
    vals = [base] + [float(np.mean([r["loss_kW"] for r in sys_rows if r["algorithm"] == a])) for a in algos]

    plt.figure(figsize=(8, 5))
    plt.bar(labels, vals)
    plt.ylabel("Power loss (kW)")
    plt.title(f"Power Loss Comparison - {system}")
    plt.xticks(rotation=15)
    plt.grid(axis="y", alpha=0.3)

    if system == "IEEE33":
        _save_plot(out_dir / "Fig_PowerLoss_Comparison_IEEE33.png", out_dir / "Fig_PowerLoss_Comparison_IEEE33.pdf")


def fig_voltage_profile(rows: List[Dict[str, Any]], meta: Dict[str, Any], system: str, out_dir: Path) -> None:
    import matplotlib.pyplot as plt

    _update_mpl_style()

    sys_rows = [r for r in rows if r["system"] == system]
    if not sys_rows:
        return
    base = meta["systems_meta"][system]["baseline_voltage_profile_pu"]
    cand = [r for r in sys_rows if r["algorithm"] == "GA-PSO"]
    if not cand:
        cand = sys_rows
    best = min(cand, key=lambda x: float(x["fitness"]))
    prof = best.get("voltage_profile_pu", [])

    x = np.arange(1, len(base) + 1)
    plt.figure(figsize=(8, 5))
    plt.plot(x, base, label="baseline", linewidth=1.8)
    plt.plot(x, prof, label=f"best_{best['algorithm']}", linewidth=1.8)
    plt.axhline(0.95, linestyle="--", linewidth=1.0)
    plt.axhline(1.05, linestyle="--", linewidth=1.0)
    plt.xlabel("Node")
    plt.ylabel("Voltage (p.u.)")
    plt.title(f"Voltage Profile - {system}")
    plt.grid(alpha=0.3)
    plt.legend()

    if system == "IEEE33":
        _save_plot(out_dir / "Fig_VoltageProfile_IEEE33.png", out_dir / "Fig_VoltageProfile_IEEE33.pdf")


def fig_success_rate(rows: List[Dict[str, Any]], out_dir: Path) -> None:
    import matplotlib.pyplot as plt

    _update_mpl_style()
    systems = sorted({r["system"] for r in rows})
    algos = sorted({r["algorithm"] for r in rows})
    if "IEEE33" not in systems or not algos:
        return

    sys_rows = [r for r in rows if r["system"] == "IEEE33"]
    vals = []
    for a in algos:
        arr = [1.0 if bool(r.get("success", False)) else 0.0 for r in sys_rows if r["algorithm"] == a]
        vals.append(100.0 * float(np.mean(arr)) if arr else 0.0)

    plt.figure(figsize=(8, 5))
    plt.bar(algos, vals)
    plt.ylabel("Success rate (%)")
    plt.title("Success Rate - IEEE33")
    plt.ylim(0, 100)
    plt.grid(axis="y", alpha=0.3)
    plt.xticks(rotation=15)
    _save_plot(out_dir / "Fig_SuccessRate_IEEE33.png", out_dir / "Fig_SuccessRate_IEEE33.pdf")


def fig_heatmap_metric(rows: List[Dict[str, Any]], metric: str, out_dir: Path) -> None:
    import matplotlib.pyplot as plt

    _update_mpl_style()
    systems = sorted({r["system"] for r in rows})
    algos = sorted({r["algorithm"] for r in rows})
    if not systems or not algos:
        return

    m = np.zeros((len(systems), len(algos)), dtype=float)
    for i, s in enumerate(systems):
        for j, a in enumerate(algos):
            vals = [float(r[metric]) for r in rows if r["system"] == s and r["algorithm"] == a]
            m[i, j] = float(np.mean(vals)) if vals else np.nan

    plt.figure(figsize=(8, 4.8))
    im = plt.imshow(m, aspect="auto")
    plt.colorbar(im)
    plt.xticks(np.arange(len(algos)), algos, rotation=20)
    plt.yticks(np.arange(len(systems)), systems)
    plt.title(f"Heatmap Mean {metric}")
    _save_plot(out_dir / f"Fig_Heatmap_{metric}.png", out_dir / f"Fig_Heatmap_{metric}.pdf")


def fig_parallel_coordinates(rows: List[Dict[str, Any]], system: str, out_dir: Path) -> None:
    import matplotlib.pyplot as plt

    _update_mpl_style()
    sys_rows = [r for r in rows if r["system"] == system]
    if not sys_rows:
        return
    algos = sorted({r["algorithm"] for r in sys_rows})
    metrics = ["fitness", "loss_kW", "voltage_dev_pu", "cost_abs_usd_per_year"]

    data = []
    for a in algos:
        vals = []
        for m in metrics:
            arr = [float(r[m]) for r in sys_rows if r["algorithm"] == a]
            vals.append(float(np.mean(arr)) if arr else 0.0)
        data.append(vals)
    data = np.array(data, dtype=float)
    mins = np.min(data, axis=0)
    maxs = np.max(data, axis=0)
    denom = np.where(maxs - mins < 1e-12, 1.0, maxs - mins)
    norm = (data - mins) / denom

    plt.figure(figsize=(8, 5))
    x = np.arange(len(metrics))
    for i, a in enumerate(algos):
        plt.plot(x, norm[i], marker="o", label=a)
    plt.xticks(x, metrics)
    plt.ylim(0, 1)
    plt.grid(alpha=0.3)
    plt.title(f"Parallel Coordinates - {system}")
    plt.legend(ncol=2)
    _save_plot(out_dir / f"Fig_ParallelCoordinates_{system}.png", out_dir / f"Fig_ParallelCoordinates_{system}.pdf")


def fig_pareto(rows: List[Dict[str, Any]], system: str, out_dir: Path) -> None:
    import matplotlib.pyplot as plt

    _update_mpl_style()
    cand = [r for r in rows if r["system"] == system and r["algorithm"] in {"NSGA-III", "MOEA/D"}]
    if not cand:
        return

    x = [float(r["loss_kW"]) for r in cand]
    y = [float(r["voltage_dev_pu"]) for r in cand]
    c = [float(r["cost_abs_usd_per_year"]) for r in cand]

    plt.figure(figsize=(8, 5))
    sc = plt.scatter(x, y, c=c, cmap="viridis", alpha=0.85)
    plt.colorbar(sc, label="cost_abs_usd_per_year")
    plt.xlabel("loss_kW")
    plt.ylabel("voltage_dev_pu")
    plt.title(f"Pareto Front Approximation - {system}")
    plt.grid(alpha=0.3)
    _save_plot(out_dir / f"Fig_ParetoFront_{system}.png", out_dir / f"Fig_ParetoFront_{system}.pdf")


def generate_tables(summary: List[Dict[str, Any]], meta: Dict[str, Any], stat_csv: Path, tables_dir: Path) -> None:
    _ensure(tables_dir)

    # Table_A
    table_a_csv = tables_dir / "Table_A_System_and_Config.csv"
    with table_a_csv.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(
            [
                "system",
                "num_buses",
                "num_lines",
                "num_capacitors",
                "num_der",
                "num_bess",
                "alpha",
                "beta",
                "gamma",
                "voltage_limits",
                "power_flow_model",
                "max_iterations",
            ]
        )
        for s, m in meta["systems_meta"].items():
            writer.writerow(
                [
                    s,
                    m["num_buses"],
                    m["num_lines"],
                    m["num_capacitors"],
                    m["num_der"],
                    m["num_bess"],
                    meta["weights"]["alpha"],
                    meta["weights"]["beta"],
                    meta["weights"]["gamma"],
                    m["voltage_limits_pu"],
                    "SurrogatePF+consistency-check",
                    meta.get("config", {}).get("iterations", "N/A"),
                ]
            )

    # Table_B
    table_b_csv = tables_dir / "Table_B_Algorithm_Settings.csv"
    algorithms = meta["algorithms"]
    with table_b_csv.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["algorithm", "population", "iterations", "budget_note"])
        for a in algorithms:
            writer.writerow([a, meta.get("config", {}).get("population", 24), meta.get("config", {}).get("iterations", 55), "Unified PF-call budget"])

    # Table_C
    table_c_csv = tables_dir / "Table_C_Performance_Summary.csv"
    with table_c_csv.open("w", newline="", encoding="utf-8") as f:
        if summary:
            keys = [
                "system",
                "algorithm",
                "num_runs",
                "mean_fitness",
                "std_fitness",
                "median_fitness",
                "best_fitness",
                "worst_fitness",
                "mean_loss_kW",
                "std_loss_kW",
                "mean_voltage_dev_pu",
                "std_voltage_dev_pu",
                "mean_cost_abs_usd_per_year",
                "std_cost_abs_usd_per_year",
                "mean_runtime_s",
                "mean_pf_calls",
                "success_rate",
            ]
            writer = csv.DictWriter(f, fieldnames=keys)
            writer.writeheader()
            for r in summary:
                writer.writerow({k: r.get(k, "") for k in keys})
        else:
            writer = csv.writer(f)
            writer.writerow(["empty"])

    # Table_D
    table_d_csv = tables_dir / "Table_D_Statistical_Tests.csv"
    if stat_csv.exists():
        _copy_file(stat_csv, table_d_csv)
    else:
        with table_d_csv.open("w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["empty"])

    # 图表字段增强：保留 success/failure 相关列

    # Word
    doc = Document()
    doc.add_heading("Paper Tables for needs0301", level=1)

    def add_csv_table(title: str, csv_path: Path):
        doc.add_heading(title, level=2)
        with csv_path.open("r", encoding="utf-8") as f:
            rows = list(csv.reader(f))
        if not rows:
            doc.add_paragraph("empty")
            return
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                table.rows[i].cells[j].text = str(val)

    add_csv_table("Table_A_System_and_Config", table_a_csv)
    add_csv_table("Table_B_Algorithm_Settings", table_b_csv)
    add_csv_table("Table_C_Performance_Summary", table_c_csv)
    add_csv_table("Table_D_Statistical_Tests", table_d_csv)

    doc.save(tables_dir / "Paper_Tables_needs0301.docx")


def create_manifest(needs_dir: Path, meta: Dict[str, Any], key_files: List[Path]) -> None:
    manifest = {
        "experiment_name": meta.get("experiment_name", "unknown"),
        "generated_at": datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z"),
        "git_commit": _get_git_commit(),
        "config_summary": {
            "weights": meta.get("weights", {}),
            "normalization": meta.get("normalization", {}),
            "currency": meta.get("currency", "USD/year"),
        },
        "systems": meta.get("systems", []),
        "algorithms": meta.get("algorithms", []),
        "num_runs": meta.get("runs_by_system", {}),
        "base_seed": meta.get("base_seed", None),
        "baseline": {
            s: {
                "loss_base_kW": m.get("loss_base_kW"),
                "cost_base_usd_per_year": m.get("cost_base_usd_per_year"),
            }
            for s, m in meta.get("systems_meta", {}).items()
        },
        "files": [],
    }

    for p in sorted(needs_dir.rglob("*")):
        if p.is_file():
            manifest["files"].append(
                {
                    "path": str(p.relative_to(needs_dir)).replace("\\", "/"),
                    "size": p.stat().st_size,
                    "sha256": _sha256_file(p),
                }
            )

    # 关键文件摘要
    manifest["key_files"] = []
    for k in key_files:
        if k.exists():
            manifest["key_files"].append(
                {
                    "path": str(k.relative_to(needs_dir)).replace("\\", "/"),
                    "size": k.stat().st_size,
                    "sha256": _sha256_file(k),
                }
            )

    (needs_dir / "manifest.json").write_text(json.dumps(manifest, indent=2, ensure_ascii=False), encoding="utf-8")


def create_readme(needs_dir: Path, input_dir: Path) -> None:
    text = f"""# README_NEEDS0301

## 一键复现

1. 最小验证（cache off/on）

```bash
python -m scripts.run_experiments --config configs/paper.yaml --out outputs/paper_v2_min --mode minimal
```

2. 正式重跑

```bash
python -m scripts.run_experiments --config configs/paper.yaml --out outputs/paper_v2 --mode full --cache on
```

3. 生成论文工件

```bash
python -m scripts.generate_paper_artifacts --input outputs/paper_v2 --out needs0301
```

## 输出目录说明

- `needs0301/data/`：run-level 原始数据与 summary、统计对比、绘图中间数据
- `needs0301/figures/`：论文图（PNG + PDF）
- `needs0301/tables/`：Word表格与CSV备份
- `needs0301/logs/`：统计检验、cache一致性、summary重算、运行命令与环境
- `needs0301/manifest.json`：工件清单、哈希与实验元信息

## 字段与单位

- `loss_kW` 单位：kW
- `voltage_dev_pu` 单位：p.u.
- `cost_abs_usd_per_year` 单位：USD/year
- 归一化：
  - `norm_loss = loss_kW / loss_base_kW`
  - `norm_vdev = voltage_dev_pu / vdev_base_pu`
  - `norm_cost = cost_abs_usd_per_year / cost_base_usd_per_year`

## 论文映射

- 收敛曲线：`Fig_Convergence_*`
- 箱线图：`Fig_Boxplot_*`
- 损耗对比：`Fig_PowerLoss_Comparison_IEEE33`
- 电压曲线：`Fig_VoltageProfile_IEEE33`
- 统计表：`tables/Table_D_Statistical_Tests.csv`
"""
    (needs_dir / "README_NEEDS0301.md").write_text(text, encoding="utf-8")


def run_generate(input_dir: Path, out_dir: Path) -> None:
    _ensure(out_dir)
    data_dir = out_dir / "data"
    fig_dir = out_dir / "figures"
    tbl_dir = out_dir / "tables"
    logs_dir = out_dir / "logs"
    scripts_dir = out_dir / "scripts"
    for d in [data_dir, fig_dir, tbl_dir, logs_dir, scripts_dir]:
        _ensure(d)

    src_data = input_dir / "data"
    runs_jsonl = src_data / "runs.jsonl"
    runs_csv = src_data / "runs.csv"
    summary_json = src_data / "summary.json"
    summary_csv = src_data / "summary.csv"
    seed_json = src_data / "seed_list.json"
    meta_json = src_data / "experiment_meta.json"

    # 基础数据复制
    for src in [runs_jsonl, runs_csv, summary_json, summary_csv, seed_json]:
        _copy_file(src, data_dir / src.name)

    rows = _read_jsonl(data_dir / "runs.jsonl")
    summary = _read_json(data_dir / "summary.json")
    meta = _read_json(meta_json)
    meta["config"] = _read_json(input_dir / "data" / "experiment_meta.json").get("config", {}) if meta_json.exists() else {}

    # 绘图中间数据
    create_boxplot_data(rows, data_dir / "boxplot_data.csv")
    create_convergence_json(rows, data_dir / "convergence_curves.json")
    create_power_loss_fixed(rows, meta, data_dir / "Power_Loss_Data_FIXED.csv")
    create_voltage_profile_fixed(rows, meta, data_dir / "Voltage_Profile_Data_FIXED.csv")

    # 统计检验
    stat_csv = data_dir / "Algorithm_Statistical_Comparison_FIXED.csv"
    stat_report = logs_dir / "stat_tests_report.md"
    run_stat_tests(data_dir / "runs.jsonl", stat_csv, stat_report)

    # 图
    plotting_available = True
    try:
        import matplotlib  # type: ignore # noqa: F401
    except Exception:
        plotting_available = False

    systems = _group_systems(rows)
    if plotting_available:
        for s in systems:
            fig_convergence(rows, s, fig_dir)

        if "IEEE33" in systems:
            fig_boxplot_metric(rows, "IEEE33", "fitness", "Fitness Boxplot - IEEE33", "Fig_Boxplot_Fitness_IEEE33", fig_dir)
            fig_power_loss_compare(rows, meta, "IEEE33", fig_dir)
            fig_voltage_profile(rows, meta, "IEEE33", fig_dir)
            fig_boxplot_metric(rows, "IEEE33", "loss_kW", "Loss Boxplot - IEEE33", "Fig_Boxplot_Loss_IEEE33", fig_dir)

        if "IEEE69" in systems:
            fig_boxplot_metric(rows, "IEEE69", "fitness", "Fitness Boxplot - IEEE69", "Fig_Boxplot_Extended_IEEE69", fig_dir)
        if "BaituF8" in systems:
            fig_boxplot_metric(rows, "BaituF8", "fitness", "Fitness Boxplot - BaituF8", "Fig_Boxplot_BaituF8", fig_dir)

        # 附加论文图
        fig_success_rate(rows, fig_dir)
        fig_heatmap_metric(rows, "fitness", fig_dir)
        fig_heatmap_metric(rows, "loss_kW", fig_dir)
        for s in systems:
            fig_parallel_coordinates(rows, s, fig_dir)
            fig_pareto(rows, s, fig_dir)
    else:
        (logs_dir / "figure_generation_report.md").write_text(
            "matplotlib 不可用，已跳过图像生成。请安装 matplotlib 后重跑 generate_paper_artifacts。",
            encoding="utf-8",
        )

    # 表
    generate_tables(summary, meta, stat_csv, tbl_dir)

    # 日志文件补充
    # cache/recompute 报告从 input logs 复制
    for name in ["cache_validation_report.md", "recompute_summary_report.md", "run_command.txt", "environment.txt"]:
        src = input_dir / "logs" / name
        if src.exists():
            _copy_file(src, logs_dir / name)
    # 最小验证中的 cache 报告补充
    min_cache = Path("outputs/paper_v2_min/logs/cache_validation_report.md")
    if min_cache.exists():
        _copy_file(min_cache, logs_dir / "cache_validation_report.md")

    # scripts 备份
    for name in ["run_experiments.py", "generate_paper_artifacts.py", "stat_tests.py", "recompute_summary.py"]:
        src = Path("scripts") / name
        if src.exists():
            _copy_file(src, scripts_dir / name)

    create_readme(out_dir, input_dir)

    key_files = [
        out_dir / "data" / "runs.csv",
        out_dir / "data" / "summary.csv",
        out_dir / "tables" / "Paper_Tables_needs0301.docx",
        out_dir / "figures" / "Fig_Convergence_IEEE33.png",
        out_dir / "figures" / "Fig_Boxplot_Fitness_IEEE33.png",
    ]
    create_manifest(out_dir, meta, key_files)


def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(description="生成论文工件")
    p.add_argument("--input", required=True, help="实验输出目录，例如 outputs/paper_v2")
    p.add_argument("--out", required=True, help="工件输出目录，例如 needs0301")
    return p


def main() -> None:
    parser = build_parser()
    args = parser.parse_args()
    run_generate(Path(args.input), Path(args.out))


if __name__ == "__main__":
    main()
